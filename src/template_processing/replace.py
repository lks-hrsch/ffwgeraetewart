import os

import docx
from docxcompose.composer import Composer

from src.logic.pathes import main_path, out_path

ALL_DOCUMENTS = {
    "Arbeitskleidung": os.path.join(main_path, "templates", "template_arbeitskleidung.docx"),
    "Einsatzkleidung": os.path.join(main_path, "templates", "template_einsatzkleidung.docx"),
    "Handschuhe": os.path.join(main_path, "templates", "template_handschuhe.docx"),
    "Helm": os.path.join(main_path, "templates", "template_helm.docx"),
    "Kopfschutzhaube": os.path.join(main_path, "templates", "template_kopfschutzhaube.docx"),
    "Schuhe": os.path.join(main_path, "templates", "template_schuhe.docx"),
}


def replace(object: docx.text, search: str, replace: str):
    if search in object.text:
        object.text = object.text.replace(search, replace)


def find_and_replace(document: docx.Document, parameters: dict) -> docx.Document:
    for tables in document.tables:
        for rows in tables.rows:
            for cell in rows.cells:
                for paragraph in cell.paragraphs:
                    for s, r in parameters.items():
                        replace(paragraph, s, r if r is not None else "")
    return document


def compose_specificpsa_for_member(parameters: dict, cloth: str):
    doc = docx.Document(ALL_DOCUMENTS[cloth])
    find_and_replace(doc, parameters)
    composed_master = Composer(doc)
    composed_master.save(out_path)


def compose_specificpsa_with_path(parameters: dict, file_name: str):
    file_path = os.path.join(main_path, *file_name.split("/"))
    doc = docx.Document(file_path)
    find_and_replace(doc, parameters)
    composed_master = Composer(doc)
    composed_master.save(out_path)


def compose_wholepsa_for_member(parameters: dict):
    arbeitskleidung = docx.Document(ALL_DOCUMENTS["Arbeitskleidung"])
    find_and_replace(arbeitskleidung, parameters)
    arbeitskleidung.add_page_break()
    composed_master = Composer(arbeitskleidung)

    einsatzkleidung = docx.Document(ALL_DOCUMENTS["Einsatzkleidung"])
    find_and_replace(einsatzkleidung, parameters)
    einsatzkleidung.add_page_break()
    composed_master.append(einsatzkleidung)

    handschuhe = docx.Document(ALL_DOCUMENTS["Handschuhe"])
    find_and_replace(handschuhe, parameters)
    handschuhe.add_page_break()
    composed_master.append(handschuhe)

    helm = docx.Document(ALL_DOCUMENTS["Helm"])
    find_and_replace(helm, parameters)
    helm.add_page_break()
    composed_master.append(helm)

    kopfschutzhaube = docx.Document(ALL_DOCUMENTS["Kopfschutzhaube"])
    find_and_replace(kopfschutzhaube, parameters)
    kopfschutzhaube.add_page_break()
    composed_master.append(kopfschutzhaube)

    schuhe = docx.Document(ALL_DOCUMENTS["Schuhe"])
    find_and_replace(schuhe, parameters)
    composed_master.append(schuhe)

    composed_master.save(out_path)


def compose_wholepsa(parameters):
    length: int = len(parameters)

    compose_wholepsa_for_member(parameters[0])
    master = docx.Document(out_path)
    if length > 1:
        master.add_page_break()
    composed_master = Composer(master)

    for x in range(1, length):
        compose_wholepsa_for_member(parameters[x])
        new_file = docx.Document(out_path)
        if x != length - 1:
            new_file.add_page_break()
        composed_master.append(new_file)

    composed_master.save(out_path)
