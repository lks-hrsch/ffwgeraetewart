import math
import os
import re

import docx
from docxcompose.composer import Composer

from src.logic.pathes import main_path, out_path

from .paragraph_replace_text import paragraph_replace_text

ALL_DOCUMENTS = {
    "BaseEquipment": os.path.join(main_path, "templates", "template_equipment.docx"),
}


def find_and_replace(document: docx.Document, parameters: dict) -> docx.Document:
    for table in document.tables:
        tmpStyle = table.style
        for rows in table.rows:
            for cell in rows.cells:
                for paragraph in cell.paragraphs:
                    for search, r in parameters.items():
                        paragraph_replace_text(paragraph, re.compile(search), r)
        table.style = tmpStyle
    return document


def compose_single_equipment(parameters: dict):
    doc = docx.Document(ALL_DOCUMENTS["BaseEquipment"])
    find_and_replace(doc, parameters)

    checkRow = ["test_date", "remark", "testVision", "testFunction", "tester"]
    if len(parameters["checks"]) <= 9:
        find_and_replace(doc, {"pagenumber": "1"})
        parameterChecks: dict = {}
        for iterChecks in range(1, 10):
            for check, iter in zip(checkRow, range(0, 5)):
                try:
                    parameterChecks[check + str(iterChecks)] = parameters["checks"][iterChecks - 1][iter]
                except IndexError:
                    parameterChecks[check + str(iterChecks)] = ""
        find_and_replace(doc, parameterChecks)
        composed_master = Composer(doc)
    else:
        find_and_replace(doc, {"pagenumber": str(1)})
        parameterChecks: dict = {}
        for iterChecks in range(1, 10):
            for check, iter in zip(checkRow, range(0, 5)):
                try:
                    parameterChecks[check + str(iterChecks)] = parameters["checks"][((1 - 1) * 9) + (iterChecks - 1)][
                        iter
                    ]
                except IndexError:
                    parameterChecks[check + str(iterChecks)] = ""
        find_and_replace(doc, parameterChecks)
        composed_master = Composer(doc)

        for iterPages in range(2, math.ceil(len(parameters["checks"]) / 9) + 1):
            doc = docx.Document(ALL_DOCUMENTS["BaseEquipment"])
            find_and_replace(doc, parameters)
            find_and_replace(doc, {"pagenumber": str(iterPages)})
            parameterChecks: dict = {}
            for iterChecks in range(1, 10):
                for check, iter in zip(checkRow, range(0, 5)):
                    try:
                        parameterChecks[check + str(iterChecks)] = parameters["checks"][
                            ((iterPages - 1) * 9) + (iterChecks - 1)
                        ][iter]
                    except IndexError:
                        parameterChecks[check + str(iterChecks)] = ""
            find_and_replace(doc, parameterChecks)
            composed_master.append(doc)

    composed_master.save(out_path)
    return composed_master


def compose_multiple_equipment(parametersList: list):
    first = True
    master = None
    for parameters in parametersList:
        if first:
            master = compose_single_equipment(parameters)
            first = False
        else:
            master.append(compose_single_equipment(parameters).doc)
    master.save(out_path)
    pass
